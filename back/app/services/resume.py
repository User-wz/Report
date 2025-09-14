from docx import Document
import os
from enum import Enum
from docx.oxml.xmlchemy import BaseOxmlElement


class ContentType(Enum):
    PARAGRAPH = 1
    TABLE = 2
    IMAGE = 3
    HEADER = 4
    FOOTER = 5



def read_word_docx_in_order(file_path):

    doc = Document(file_path)
    ordered_content = []
    image_counter = 0
    image_dir = os.path.join(os.path.dirname(file_path), 'extracted_images')
    os.makedirs(image_dir, exist_ok=True)

    # 先处理页眉
    for section in doc.sections:
        for header in section.header.paragraphs:
            if header.text.strip():
                ordered_content.append({
                    'type': ContentType.HEADER,
                    'text': header.text,
                    'style': header.style.name
                })

    # 处理文档主体内容
    for block in doc.element.body.xpath('*'):
        # 处理段落
        if block.tag.endswith('p'):
            # 找到对应的Paragraph对象
            for para in doc.paragraphs:
                if para._p is block:  # 比较内部XML元素
                    if para.text.strip():
                        ordered_content.append({
                            'type': ContentType.PARAGRAPH,
                            'text': para.text,
                            'style': para.style.name
                        })
                    break

        # 处理表格
        elif block.tag.endswith('tbl'):
            # 找到对应的Table对象
            for table in doc.tables:
                if table._tbl is block:  # 比较内部XML元素
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text)
                        table_data.append(row_data)
                    ordered_content.append({
                        'type': ContentType.TABLE,
                        'data': table_data
                    })
                    break

    # 处理图片
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_ext = rel.target_ref.split(".")[-1]
            image_path = os.path.join(image_dir, f"image_{image_counter}.{image_ext}")
            with open(image_path, 'wb') as f:
                f.write(image_data)
            ordered_content.append({
                'type': ContentType.IMAGE,
                'path': image_path,
                'position': image_counter
            })
            image_counter += 1

    # 处理页脚
    for section in doc.sections:
        for footer in section.footer.paragraphs:
            if footer.text.strip():
                ordered_content.append({
                    'type': ContentType.FOOTER,
                    'text': footer.text,
                    'style': footer.style.name
                })

    return ordered_content


def print_ordered_content(content):
    str_=""
    for item in content:
        if item['type'] == ContentType.PARAGRAPH:
            print()
            str_+=f"[段落] {item['text']} (样式: {item['style']})"
        elif item['type'] == ContentType.TABLE:
            d="[表格]"
            print("[表格]")
            for row in item['data']:
                d+=" | ".join(row)
                print(" | ".join(row))
            str_+=d
        elif item['type'] == ContentType.IMAGE:
            print(f"[图片] 保存路径: {item['path']}")
            str_+=f"[图片] 保存路径: {item['path']}"
        elif item['type'] == ContentType.HEADER:
            print(f"[页眉] {item['text']}")
            str_+=f"[页眉] {item['text']}"
        elif item['type'] == ContentType.FOOTER:
            print(f"[页脚] {item['text']}")
            str_+=f"[页脚] {item['text']}"
    return str_



if __name__ == '__main__':
    # 使用示例
    # ordered_data = read_word_docx_in_order('resume.docx')
    # print_ordered_content(ordered_data)
    ordered_data = read_word_docx_in_order(r'C:\Users\Administrator\Desktop\人\王震-简历 - python.docx')
    # print(ordered_data)
    d=print_ordered_content(ordered_data)
    print("+"*30)
    print(d)